from pathlib import Path
from docx import Document

from app.parsers.base import BaseParser


class DOCXParser(BaseParser):
    """DOCX document parser using python-docx."""

    def supports(self, file_type: str) -> bool:
        return file_type.lower() == "docx"

    def parse(self, file_path: Path) -> str:
        """Extract text from DOCX document."""
        doc = Document(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n\n".join(text_parts)

    def get_metadata(self, file_path: Path) -> dict:
        """Extract DOCX metadata."""
        metadata = super().get_metadata(file_path)

        try:
            doc = Document(file_path)
            core_props = doc.core_properties

            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author

            metadata["paragraph_count"] = len(doc.paragraphs)
        except Exception:
            pass

        return metadata
